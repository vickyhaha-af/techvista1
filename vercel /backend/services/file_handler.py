"""
File handler service — extracts raw text from PDF and DOCX files locally.
No AI API calls needed for text extraction.
"""
import io
import fitz  # PyMuPDF
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    text = ""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    if not text.strip():
        raise ValueError("PDF appears to be empty or contains only images (scanned PDFs are not supported in MVP)")
    
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    if not text.strip():
        raise ValueError("DOCX appears to be empty")
    
    return text.strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from a file based on its extension."""
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {filename}. Only PDF and DOCX are supported.")
